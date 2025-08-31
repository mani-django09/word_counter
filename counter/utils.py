import re
import math
from collections import Counter
from typing import Dict, List, Tuple
import docx
import fitz  # PyMuPDF
from django.conf import settings

class TextAnalyzer:
    STOPWORDS = {
        'a', 'an', 'and', 'are', 'as', 'at', 'be', 'by', 'for', 'from',
        'has', 'he', 'in', 'is', 'it', 'its', 'of', 'on', 'that', 'the',
        'to', 'was', 'were', 'will', 'with', 'the', 'this', 'but', 'they',
        'have', 'had', 'what', 'said', 'each', 'which', 'she', 'do', 'how',
        'their', 'if', 'up', 'out', 'many', 'then', 'them', 'these', 'so',
        'some', 'her', 'would', 'make', 'like', 'into', 'him', 'time', 'two',
        'more', 'go', 'no', 'way', 'could', 'my', 'than', 'first', 'been',
        'call', 'who', 'its', 'now', 'find', 'long', 'down', 'day', 'did',
        'get', 'come', 'made', 'may', 'part'
    }
    
    def __init__(self, text: str):
        self.text = text
        self._words = None
        self._sentences = None
        self._paragraphs = None
    
    @property
    def words(self) -> List[str]:
        if self._words is None:
            # Split by whitespace and punctuation, filter empty strings
            self._words = [w for w in re.findall(r'\b\w+\b', self.text.lower()) if w]
        return self._words
    
    @property
    def sentences(self) -> List[str]:
        if self._sentences is None:
            # Split by sentence endings, filter empty strings
            sentences = re.split(r'[.!?]+', self.text)
            self._sentences = [s.strip() for s in sentences if s.strip()]
        return self._sentences
    
    @property
    def paragraphs(self) -> List[str]:
        if self._paragraphs is None:
            # Split by double newlines or more, filter empty strings
            paragraphs = re.split(r'\n\s*\n', self.text.strip())
            self._paragraphs = [p.strip() for p in paragraphs if p.strip()]
        return self._paragraphs
    
    def get_stats(self) -> Dict:
        words = self.words
        
        return {
            'word_count': len(words),
            'character_count': len(self.text),
            'character_count_no_spaces': len(self.text.replace(' ', '')),
            'sentence_count': len(self.sentences),
            'paragraph_count': len(self.paragraphs),
            'reading_time': self.calculate_reading_time(len(words)),
            'speaking_time': self.calculate_speaking_time(len(words)),
            'keyword_density': self.get_keyword_density(words)
        }
    
    def calculate_reading_time(self, word_count: int) -> Dict:
        # Average reading speed: 200-250 WPM
        wpm = 225
        minutes = math.ceil(word_count / wpm)
        return {
            'minutes': minutes,
            'seconds': math.ceil((word_count / wpm) * 60) % 60 if minutes == 0 else 0
        }
    
    def calculate_speaking_time(self, word_count: int) -> Dict:
        # Average speaking speed: 150-160 WPM
        wpm = 155
        minutes = math.ceil(word_count / wpm)
        return {
            'minutes': minutes,
            'seconds': math.ceil((word_count / wpm) * 60) % 60 if minutes == 0 else 0
        }
    
    def get_keyword_density(self, words: List[str]) -> List[Dict]:
        if not words:
            return []
        
        # Filter words: remove stopwords, words < 3 chars
        filtered_words = [
            word for word in words 
            if len(word) >= 3 and word not in self.STOPWORDS
        ]
        
        if not filtered_words:
            return []
        
        # Count occurrences
        word_counts = Counter(filtered_words)
        total_words = len(filtered_words)
        
        # Calculate density and return top 10
        keyword_data = []
        for word, count in word_counts.most_common(10):
            density = (count / total_words) * 100
            keyword_data.append({
                'word': word,
                'count': count,
                'density': round(density, 2)
            })
        
        return keyword_data

class FileProcessor:
    @staticmethod
    def extract_text_from_docx(file) -> str:
        try:
            doc = docx.Document(file)
            text = []
            for paragraph in doc.paragraphs:
                text.append(paragraph.text)
            return '\n'.join(text)
        except Exception as e:
            raise ValueError(f"Error processing DOCX file: {str(e)}")
    
    @staticmethod
    def extract_text_from_pdf(file) -> str:
        if not settings.ENABLE_PDF_SUPPORT:
            raise ValueError("PDF support is not enabled")
        
        try:
            doc = fitz.open(stream=file.read(), filetype="pdf")
            text = []
            for page in doc:
                text.append(page.get_text())
            doc.close()
            return '\n'.join(text)
        except Exception as e:
            raise ValueError(f"Error processing PDF file: {str(e)}")
    
    @staticmethod
    def extract_text_from_txt(file) -> str:
        try:
            return file.read().decode('utf-8')
        except UnicodeDecodeError:
            try:
                file.seek(0)
                return file.read().decode('latin-1')
            except Exception as e:
                raise ValueError(f"Error processing text file: {str(e)}")

class TextCleaner:
    @staticmethod
    def remove_extra_spaces(text: str) -> str:
        # Remove multiple spaces, tabs, and normalize line breaks
        text = re.sub(r'[ \t]+', ' ', text)  # Multiple spaces/tabs to single space
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)  # Multiple line breaks to double
        text = re.sub(r'^\s+|\s+$', '', text, flags=re.MULTILINE)  # Trim lines
        return text.strip()

class CaseConverter:
    @staticmethod
    def to_upper(text: str) -> str:
        return text.upper()
    
    @staticmethod
    def to_lower(text: str) -> str:
        return text.lower()
    
    @staticmethod
    def to_title(text: str) -> str:
        return text.title()
    
    @staticmethod
    def to_sentence(text: str) -> str:
        if not text:
            return text
        
        # Split by sentence endings and capitalize first letter of each sentence
        sentences = re.split(r'([.!?]+)', text)
        result = []
        
        for i, part in enumerate(sentences):
            if i % 2 == 0:  # Text part, not punctuation
                part = part.strip()
                if part:
                    part = part[0].upper() + part[1:].lower() if len(part) > 1 else part.upper()
                result.append(part)
            else:  # Punctuation part
                result.append(part)
        
        return ''.join(result)